"""
requirements: requests, pandas, python-docx, openpyxl, pydantic
"""

import os
import requests
import pandas as pd
from docx import Document
import logging
from typing import List, Union, Generator, Iterator
from pydantic import BaseModel

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class Pipeline:

    class Valves(BaseModel):
        # Параметры пайплайна, задаваемые внешне
        mistral_api_key: str
        mistral_model: str = "mistral-small-latest"

    def __init__(self, api_key: str):
        self.name = "Нормативные документы из DOCX в Excel"
        self.valves = self.Valves(mistral_api_key=api_key)

    def extract_text_from_docx(self, path: str) -> str:
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs]
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    texts.append(cell.text)
        return "\n".join(texts)

    def query_model(self, prompt: str) -> str:
        API_URL = "https://api.mistral.ai/v1/chat/completions"
        HEADERS = {
            "Authorization": f"Bearer {self.valves.mistral_api_key}",
            "Content-Type": "application/json",
        }

        system_message = {
            "role": "system",
            "content": (
                "You are a smart assistant that analyzes user requests carefully and answers in Russian. "
                "Найди в тексте все нормативные документы (ГОСТ, Приказ, Постановление). "
                "Для каждого выведи в формате: Тип документа; Номер и дата; Статус (действует/не действует/статус неизвестен). "
                "Отвечай построчно, каждую запись на новой строке, поля разделяй точкой с запятой ';'. "
                "Пример:\n"
                "ГОСТ; ГОСТ 1234-56 от 01.01.2000; действует\n"
                "Приказ; Приказ №12 от 05.05.2022; статус неизвестен\n"
            )
        }
        user_message = {"role": "user", "content": prompt}

        data = {
            "model": self.valves.mistral_model,
            "messages": [system_message, user_message],
            "temperature": 0.1,
        }

        response = requests.post(API_URL, headers=HEADERS, json=data)
        if response.status_code != 200:
            logger.error(f"Ошибка API Mistral: {response.status_code} - {response.text}")
            return ""

        content = response.json()["choices"][0]["message"]["content"]
        return content

    def parse_text_to_dataframe(self, text: str) -> pd.DataFrame:
        rows = []
        for line in text.strip().splitlines():
            if not line.strip():
                continue
            parts = [part.strip() for part in line.split(";")]
            if len(parts) != 3:
                logger.warning(f"Строка пропущена (не 3 поля): {line}")
                continue
            rows.append({
                "Тип документа": parts[0],
                "Номер и дата": parts[1],
                "Статус": parts[2],
            })
        return pd.DataFrame(rows)

    def save_to_excel(self, df: pd.DataFrame, output_path: str) -> bool:
        try:
            df.to_excel(output_path, index=False)
            logger.info(f"Отчёт сохранён в {output_path}")
            return True
        except Exception as e:
            logger.error(f"Ошибка при сохранении Excel: {e}")
            return False

    def pipe(
        self,
        user_message: dict,
        model_id: str,
        messages: List[dict],
        body: dict
    ) -> Union[str, Generator, Iterator]:

        """
        user_message: dict - сообщение пользователя
        model_id: str - id модели (например, "mistral-small-latest")
        messages: List[dict] - список сообщений (игнорируется, т.к. формируем сами)
        body: dict - словарь с параметрами, ожидается ключ "file_path" - путь к docx файлу
        """

        file_path = body.get("file_path", "")
        if not file_path or not os.path.isfile(file_path):
            return "Файл не найден или путь не указан."

        # Извлечь текст из docx
        text = self.extract_text_from_docx(file_path)

        prompt = f"""
В тексте ниже найди все упоминания нормативных документов: ГОСТ, Приказ, Постановление.
Для каждого документа укажи:
- Тип документа
- Номер и дату
- Статус действия документа (действует, не действует, статус неизвестен)

Текст для анализа:
'''{text}'''
"""

        # Запрос к модели
        response_text = self.query_model(prompt)
        if not response_text:
            return "Ошибка при получении ответа от модели."

        # Парсинг в DataFrame
        df = self.parse_text_to_dataframe(response_text)
        if df.empty:
            return "Не удалось распарсить ответ модели или документов не найдено."

        output_path = file_path.rsplit(".", 1)[0] + "_отчет.xlsx"
        success = self.save_to_excel(df, output_path)
        if not success:
            return "Ошибка при сохранении отчёта."

        return f"Обработка завершена. Отчёт сохранён: {output_path}"


if __name__ == "__main__":
    # Пример запуска пайплайна
    api_key = "arA3rgXk1YQskcsxcvcUYHtsEPhC3plI"  # Твой ключ
    pipeline = Pipeline(api_key=api_key)

    file_path = input("Введите путь к документу .docx: ").strip()

    result = pipeline.pipe(
        user_message={},  # В данном примере не используется
        model_id=pipeline.valves.mistral_model,
        messages=[],
        body={"file_path": file_path}
    )
    print(result)