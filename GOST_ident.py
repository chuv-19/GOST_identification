import os
import requests
from docx import Document
import logging
import pandas as pd

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

model = "mistral-small-latest"
key = "arA3rgXk1YQskcsxcvcUYHtsEPhC3plI"
API_URL = "https://api.mistral.ai/v1/chat/completions"
HEADERS = {
    "Authorization": f"Bearer {key}",
    "Content-Type": "application/json",
}

def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    texts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return "\n".join(texts)

def query_model(prompt: str) -> str:
    system_message = {
        "role": "system",
        "content": (
            "You are a smart assistant that analyzes user requests carefully and answers in Russian. "
            "Найди в тексте все нормативные документы (ГОСТ, Приказ, Постановление). "
            "Для каждого выведи в формате: Тип документа; Номер и дата; Статус (действует/не действует/статус неизвестен). "
            "Отвечай построчно, каждую запись на новой строке, поля разделяй точкой с запятой ';'. "
            "Пример:\n"
            "ГОСТ; ГОСТ 1234-56 от 01.01.2000; действует\n"
            "Приказ; Приказ №12 от 05.05.2022; статус неизвестен\n"
        )
    }
    user_message = {"role": "user", "content": prompt}

    data = {
        "model": model,
        "messages": [system_message, user_message],
        "temperature": 0.1,
    }

    response = requests.post(API_URL, headers=HEADERS, json=data)
    if response.status_code != 200:
        logger.error(f"Ошибка API Mistral: {response.status_code} - {response.text}")
        return ""

    content = response.json()["choices"][0]["message"]["content"]
    return content

def parse_text_to_dataframe(text: str):
    rows = []
    for line in text.strip().splitlines():
        if not line.strip():
            continue
        parts = [part.strip() for part in line.split(";")]
        if len(parts) != 3:
            logger.warning(f"Строка пропущена (не 3 поля): {line}")
            continue
        rows.append({
            "Тип документа": parts[0],
            "Номер и дата": parts[1],
            "Статус": parts[2],
        })
    return pd.DataFrame(rows)

def process_docx(file_path: str) -> str:
    if not file_path or not os.path.isfile(file_path):
        return "Файл не найден или путь не указан."

    text = extract_text_from_docx(file_path)

    prompt = f"""
В тексте ниже найди все упоминания нормативных документов: ГОСТ, Приказ, Постановление.
Для каждого документа укажи:
- Тип документа
- Номер и дату
- Статус действия документа (действует, не действует, статус неизвестен)

Текст для анализа:
'''{text}'''
"""

    response_text = query_model(prompt)
    if not response_text:
        return "Ошибка при получении ответа от модели."

    df = parse_text_to_dataframe(response_text)
    if df.empty:
        return "Не удалось распарсить ответ модели или документов не найдено."

    logger.info(f"Количество записей для сохранения: {len(df)}")
    print(df)

    output_path = file_path.rsplit(".", 1)[0] + "_отчет.xlsx"
    try:
        df.to_excel(output_path, index=False)
        logger.info(f"Отчёт сохранён в {output_path}")
    except Exception as e:
        logger.error(f"Ошибка при сохранении Excel: {e}")
        return f"Ошибка при сохранении Excel: {e}"

    return f"Обработка завершена. Отчёт сохранён: {output_path}"

if __name__ == "__main__":
    file_path = input("Введите путь к документу .docx: ").strip()
    result = process_docx(file_path)
    print(result)